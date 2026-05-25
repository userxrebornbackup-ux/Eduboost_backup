"""
etl_pipeline.py — Eduboost Document & Training Data Pipeline
=============================================================
Implements Phases 0-7 of the ETL roadmap:

  Phase 0  — Data contracts & schemas
  Phase 1  — Source registry & acquisition
  Phase 2  — Raw document storage model
  Phase 3  — Extraction layer (PDF/DOCX/TXT/HTML/MD/CSV)
  Phase 4  — Normalization & cleaning
  Phase 5  — Metadata enrichment
  Phase 6  — Chunking & segmentation
  Phase 7  — Quality validation & scoring

Storage backend: SQLite (drop-in swap for PostgreSQL via SQLAlchemy URL).
File storage: local filesystem (drop-in swap for S3 via object_store adapter).

Usage
-----
    from etl_pipeline import EduboostETL, IngestRequest

    etl = EduboostETL(db_url="sqlite:///eduboost_etl.db", storage_root="./data")
    etl.init_db()

    job = etl.ingest(IngestRequest(
        file_path="grade4_maths_textbook.pdf",
        source_type="manual_upload",
        uploaded_by="content_team",
        document_type="textbook",
        grade=4,
        subject="mathematics",
    ))
    result = etl.run_full_pipeline(job.document_id)
    print(result.quality_score)
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import sqlite3
import unicodedata
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Optional
from uuid import uuid4

# ---------------------------------------------------------------------------
# ── Optional heavy deps — graceful fallbacks for environments without them ──
# ---------------------------------------------------------------------------
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import docx as python_docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False


# ===========================================================================
# PHASE 0 — DATA CONTRACTS & SCHEMAS
# ===========================================================================

class DocumentType(str, Enum):
    textbook          = "textbook"
    workbook          = "workbook"
    teacher_guide     = "teacher_guide"
    lesson_plan       = "lesson_plan"
    assessment_rubric = "assessment_rubric"
    past_paper        = "past_paper"
    memorandum        = "memorandum"
    curriculum_statement = "curriculum_statement"
    subject_policy    = "subject_policy"
    act_regulation    = "act_regulation"
    governance_doc    = "governance_doc"
    circular          = "circular"
    remediation       = "remediation"
    enrichment        = "enrichment"
    worksheet         = "worksheet"
    study_material    = "study_material"
    generated         = "generated"
    unknown           = "unknown"


class SourceType(str, Enum):
    manual_upload      = "manual_upload"
    public_url         = "public_url"
    government_repo    = "government_repo"
    internal_upload    = "internal_upload"
    partner_content    = "partner_content"
    generated_content  = "generated_content"


class ProcessingStatus(str, Enum):
    raw               = "raw"
    acquired          = "acquired"
    extracted         = "extracted"
    normalized        = "normalized"
    metadata_enriched = "metadata_enriched"
    chunked           = "chunked"
    validated         = "validated"
    needs_review      = "needs_review"
    approved          = "approved"
    indexed           = "indexed"
    training_ready    = "training_ready"
    rejected          = "rejected"
    archived          = "archived"


class LicenseStatus(str, Enum):
    open_license    = "open_license"
    creative_commons = "creative_commons"
    government_open = "government_open"
    proprietary     = "proprietary"
    unknown         = "unknown"
    restricted      = "restricted"


class ChunkType(str, Enum):
    section    = "section"
    topic      = "topic"
    lesson     = "lesson"
    legal_clause = "legal_clause"
    assessment_question = "assessment_question"
    answer_memo = "answer_memo"
    table      = "table"
    summary    = "summary"
    glossary   = "glossary"
    paragraph  = "paragraph"


# ---------------------------------------------------------------------------
# Dataclasses (Phase 0 schemas — map 1-to-1 to DB tables)
# ---------------------------------------------------------------------------

@dataclass
class DocumentSource:
    """Phase 1: source_registry row."""
    source_id:        str
    source_type:      str
    source_url:       Optional[str]
    uploaded_by:      str
    license_status:   str
    retrieved_at:     str
    checksum:         Optional[str]  # SHA-256 of the original file
    is_trusted:       bool = True
    notes:            str = ""


@dataclass
class Document:
    """Phase 8: canonical documents row — master record."""
    document_id:      str
    source_id:        str
    title:            str
    description:      str
    document_type:    str
    subject:          Optional[str]
    grade:            Optional[int]
    phase:            Optional[str]        # CAPS phase (Foundation/Intermediate/Senior/FET)
    curriculum:       Optional[str]        # e.g. "CAPS", "IEB", "Cambridge"
    country:          str
    province:         Optional[str]
    language:         str
    publisher:        Optional[str]
    author:           Optional[str]
    publication_year: Optional[int]
    version:          str
    license_status:   str
    source_url:       Optional[str]
    checksum:         str
    file_path_raw:    str
    file_size_bytes:  int
    page_count:       int
    mime_type:        str
    processing_status: str
    quality_score:    float
    training_readiness: bool
    created_at:       str
    updated_at:       str
    rejected_reason:  Optional[str] = None
    reviewer_notes:   Optional[str] = None
    reviewed_by:      Optional[str] = None


@dataclass
class DocumentChunk:
    """Phase 6: document_chunks row."""
    chunk_id:         str
    document_id:      str
    chunk_type:       str
    chunk_index:      int
    parent_chunk_id:  Optional[str]
    heading:          str
    content:          str
    token_count:      int
    page_start:       Optional[int]
    page_end:         Optional[int]
    section_path:     str        # e.g. "Chapter 3 > Section 2 > 2.1"
    curriculum_code:  Optional[str]
    created_at:       str


@dataclass
class QualityCheckResult:
    """Phase 7: quality validation report."""
    document_id:          str
    metadata_score:       float   # 0–1
    extraction_score:     float
    structure_score:      float
    completeness_score:   float
    provenance_score:     float
    training_suitability: float
    quality_score:        float   # weighted composite
    status:               str     # validated / needs_review / rejected
    issues:               list[str] = field(default_factory=list)
    checked_at:           str = ""

    def __post_init__(self):
        if not self.checked_at:
            self.checked_at = _now()

    @staticmethod
    def compute_composite(meta, extract, struct, complete, prov, training) -> float:
        return (
            meta     * 0.20 +
            extract  * 0.20 +
            struct   * 0.20 +
            complete * 0.20 +
            prov     * 0.10 +
            training * 0.10
        )


@dataclass
class ProcessingJob:
    """Tracks one run through the pipeline."""
    job_id:       str
    document_id:  str
    stage:        str
    status:       str    # running / success / failed
    started_at:   str
    finished_at:  Optional[str] = None
    error:        Optional[str] = None
    output_path:  Optional[str] = None


@dataclass
class IngestRequest:
    """Input DTO for etl.ingest()."""
    file_path:      str
    source_type:    str = SourceType.manual_upload
    uploaded_by:    str = "system"
    document_type:  str = DocumentType.unknown
    source_url:     Optional[str] = None
    license_status: str = LicenseStatus.unknown
    grade:          Optional[int] = None
    subject:        Optional[str] = None
    language:       str = "en"
    title:          Optional[str] = None
    notes:          str = ""


# ===========================================================================
# HELPERS
# ===========================================================================

def _now() -> str:
    return datetime.now(timezone.utc).isoformat()

def _uid() -> str:
    return str(uuid4())

def _sha256(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()

def _token_count(text: str) -> int:
    """Rough token estimate (1 token ≈ 4 chars)."""
    return max(1, len(text) // 4)

CURRICULUM_CODES = re.compile(r'\b(LO\d+|AS\d+|CAPS-\w+|\d+\.\d+\.\d+)\b')
LEGAL_SECTION    = re.compile(r'\bSection\s+\d+[A-Z]?\b|\bClause\s+\d+\b|\bArticle\s+\d+\b', re.I)
HEADING_PATTERN  = re.compile(r'^(#{1,6})\s+(.+)$', re.M)
PAGE_ARTIFACT    = re.compile(r'Page\s+\d+\s+of\s+\d+', re.I)
REPEATED_SPACE   = re.compile(r'[ \t]{2,}')
OCR_ARTIFACTS    = [
    (re.compile(r'\bl\b(?=[a-z])'), 'i'),   # l→i in OCR
    (re.compile(r'\bO\b(?=\d)'), '0'),       # O→0 in OCR
    (re.compile(r'\(\s+\)'), '()'),
]

GRADE_PATTERN    = re.compile(r'\bGrade\s+(\d+)\b|\bGr\.\s*(\d+)\b', re.I)
SUBJECT_KEYWORDS = {
    "mathematics": ["math", "algebra", "geometry", "calculus", "numeracy"],
    "english":     ["english", "language arts", "literacy", "home language"],
    "science":     ["science", "biology", "chemistry", "physics", "natural"],
    "history":     ["history", "social studies", "humanities"],
    "geography":   ["geography", "maps", "environment"],
    "accounting":  ["accounting", "financial", "ledger", "balance sheet"],
    "economics":   ["economics", "market", "supply", "demand"],
    "life_sciences": ["life sciences", "biology", "organism", "cells"],
    "physical_sciences": ["physical sciences", "chemistry", "physics", "mechanics"],
}


# ===========================================================================
# PHASE 0 — DB SCHEMA
# ===========================================================================

SCHEMA_SQL = """
CREATE TABLE IF NOT EXISTS document_sources (
    source_id       TEXT PRIMARY KEY,
    source_type     TEXT NOT NULL,
    source_url      TEXT,
    uploaded_by     TEXT NOT NULL,
    license_status  TEXT NOT NULL,
    retrieved_at    TEXT NOT NULL,
    checksum        TEXT,
    is_trusted      INTEGER DEFAULT 1,
    notes           TEXT DEFAULT ''
);

CREATE TABLE IF NOT EXISTS documents (
    document_id      TEXT PRIMARY KEY,
    source_id        TEXT NOT NULL,
    title            TEXT NOT NULL,
    description      TEXT DEFAULT '',
    document_type    TEXT NOT NULL,
    subject          TEXT,
    grade            INTEGER,
    phase            TEXT,
    curriculum       TEXT,
    country          TEXT DEFAULT 'ZA',
    province         TEXT,
    language         TEXT DEFAULT 'en',
    publisher        TEXT,
    author           TEXT,
    publication_year INTEGER,
    version          TEXT DEFAULT '1.0',
    license_status   TEXT DEFAULT 'unknown',
    source_url       TEXT,
    checksum         TEXT NOT NULL,
    file_path_raw    TEXT NOT NULL,
    file_size_bytes  INTEGER DEFAULT 0,
    page_count       INTEGER DEFAULT 0,
    mime_type        TEXT DEFAULT 'application/octet-stream',
    processing_status TEXT NOT NULL DEFAULT 'raw',
    quality_score    REAL DEFAULT 0.0,
    training_readiness INTEGER DEFAULT 0,
    created_at       TEXT NOT NULL,
    updated_at       TEXT NOT NULL,
    rejected_reason  TEXT,
    reviewer_notes   TEXT,
    reviewed_by      TEXT,
    FOREIGN KEY (source_id) REFERENCES document_sources(source_id)
);

CREATE TABLE IF NOT EXISTS document_chunks (
    chunk_id         TEXT PRIMARY KEY,
    document_id      TEXT NOT NULL,
    chunk_type       TEXT NOT NULL,
    chunk_index      INTEGER NOT NULL,
    parent_chunk_id  TEXT,
    heading          TEXT DEFAULT '',
    content          TEXT NOT NULL,
    token_count      INTEGER DEFAULT 0,
    page_start       INTEGER,
    page_end         INTEGER,
    section_path     TEXT DEFAULT '',
    curriculum_code  TEXT,
    created_at       TEXT NOT NULL,
    FOREIGN KEY (document_id) REFERENCES documents(document_id)
);

CREATE TABLE IF NOT EXISTS quality_checks (
    check_id             TEXT PRIMARY KEY,
    document_id          TEXT NOT NULL,
    metadata_score       REAL DEFAULT 0.0,
    extraction_score     REAL DEFAULT 0.0,
    structure_score      REAL DEFAULT 0.0,
    completeness_score   REAL DEFAULT 0.0,
    provenance_score     REAL DEFAULT 0.0,
    training_suitability REAL DEFAULT 0.0,
    quality_score        REAL DEFAULT 0.0,
    status               TEXT NOT NULL,
    issues               TEXT DEFAULT '[]',
    checked_at           TEXT NOT NULL,
    FOREIGN KEY (document_id) REFERENCES documents(document_id)
);

CREATE TABLE IF NOT EXISTS processing_jobs (
    job_id       TEXT PRIMARY KEY,
    document_id  TEXT NOT NULL,
    stage        TEXT NOT NULL,
    status       TEXT NOT NULL,
    started_at   TEXT NOT NULL,
    finished_at  TEXT,
    error        TEXT,
    output_path  TEXT
);

CREATE TABLE IF NOT EXISTS review_tasks (
    task_id      TEXT PRIMARY KEY,
    document_id  TEXT NOT NULL,
    reason       TEXT NOT NULL,
    created_at   TEXT NOT NULL,
    resolved_at  TEXT,
    resolved_by  TEXT,
    resolution   TEXT
);

CREATE INDEX IF NOT EXISTS idx_docs_status  ON documents(processing_status);
CREATE INDEX IF NOT EXISTS idx_docs_grade   ON documents(grade);
CREATE INDEX IF NOT EXISTS idx_docs_subject ON documents(subject);
CREATE INDEX IF NOT EXISTS idx_chunks_doc   ON document_chunks(document_id);
CREATE INDEX IF NOT EXISTS idx_jobs_doc     ON processing_jobs(document_id);
"""


# ===========================================================================
# PHASE 3 — EXTRACTION LAYER
# ===========================================================================

@dataclass
class ExtractionResult:
    raw_text:       str
    pages:          list[dict]   # [{page_num, text, headings}]
    tables:         list[dict]   # [{page, headers, rows}]
    headings:       list[str]
    page_count:     int
    mime_type:      str
    ocr_confidence: Optional[float]
    extraction_ok:  bool
    error:          Optional[str] = None


class Extractor:
    """Phase 3: multi-format text extraction."""

    def extract(self, path: str) -> ExtractionResult:
        p = Path(path)
        suffix = p.suffix.lower()
        dispatch = {
            ".pdf":  self._pdf,
            ".docx": self._docx,
            ".doc":  self._docx,
            ".txt":  self._txt,
            ".md":   self._txt,
            ".html": self._html,
            ".htm":  self._html,
            ".csv":  self._csv,
            ".xlsx": self._xlsx,
            ".xls":  self._xlsx,
        }
        handler = dispatch.get(suffix, self._txt)
        try:
            return handler(path)
        except Exception as e:
            return ExtractionResult(
                raw_text="", pages=[], tables=[], headings=[],
                page_count=0, mime_type="unknown", ocr_confidence=None,
                extraction_ok=False, error=str(e)
            )

    def _pdf(self, path: str) -> ExtractionResult:
        if not HAS_PYMUPDF:
            return self._txt(path)   # fallback
        pages, headings, tables = [], [], []
        raw_parts = []
        with fitz.open(path) as doc:
            for pnum, page in enumerate(doc, 1):
                text = page.get_text("text")
                raw_parts.append(text)
                page_headings = [
                    b["lines"][0]["spans"][0]["text"]
                    for b in page.get_text("dict")["blocks"]
                    if b.get("type") == 0
                    for line in b["lines"]
                    for span in line["spans"]
                    if span.get("size", 0) > 13   # heuristic: large font = heading
                ]
                headings.extend(page_headings)
                pages.append({"page_num": pnum, "text": text, "headings": page_headings})
        raw_text = "\n".join(raw_parts)
        is_scanned = len(raw_text.strip()) < 100 and len(doc) > 0
        ocr_conf = None if not is_scanned else 0.0
        return ExtractionResult(
            raw_text=raw_text, pages=pages, tables=tables,
            headings=headings, page_count=len(pages),
            mime_type="application/pdf",
            ocr_confidence=ocr_conf, extraction_ok=True
        )

    def _docx(self, path: str) -> ExtractionResult:
        if not HAS_DOCX:
            return self._txt(path)
        doc = python_docx.Document(path)
        paras, headings, pages = [], [], []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                headings.append(para.text)
            paras.append(para.text)
        raw_text = "\n".join(paras)
        tables = []
        for tbl in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in tbl.rows]
            if rows:
                tables.append({"headers": rows[0], "rows": rows[1:]})
        pages.append({"page_num": 1, "text": raw_text, "headings": headings})
        return ExtractionResult(
            raw_text=raw_text, pages=pages, tables=tables,
            headings=headings, page_count=1,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ocr_confidence=None, extraction_ok=True
        )

    def _txt(self, path: str) -> ExtractionResult:
        with open(path, "r", errors="replace") as f:
            text = f.read()
        headings = [m.group(2) for m in HEADING_PATTERN.finditer(text)]
        pages = [{"page_num": 1, "text": text, "headings": headings}]
        mime = "text/markdown" if path.endswith(".md") else "text/plain"
        return ExtractionResult(
            raw_text=text, pages=pages, tables=[], headings=headings,
            page_count=1, mime_type=mime, ocr_confidence=None, extraction_ok=True
        )

    def _html(self, path: str) -> ExtractionResult:
        with open(path, "r", errors="replace") as f:
            html = f.read()
        if HAS_BS4:
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(separator="\n")
            headings = [t.get_text() for t in soup.find_all(["h1","h2","h3","h4"])]
        else:
            text = re.sub(r'<[^>]+>', ' ', html)
            headings = []
        pages = [{"page_num": 1, "text": text, "headings": headings}]
        return ExtractionResult(
            raw_text=text, pages=pages, tables=[], headings=headings,
            page_count=1, mime_type="text/html", ocr_confidence=None, extraction_ok=True
        )

    def _csv(self, path: str) -> ExtractionResult:
        if HAS_PANDAS:
            df = pd.read_csv(path)
            text = df.to_string(index=False)
            tables = [{"headers": list(df.columns), "rows": df.values.tolist()}]
        else:
            with open(path, "r") as f:
                text = f.read()
            tables = []
        pages = [{"page_num": 1, "text": text, "headings": []}]
        return ExtractionResult(
            raw_text=text, pages=pages, tables=tables, headings=[],
            page_count=1, mime_type="text/csv", ocr_confidence=None, extraction_ok=True
        )

    def _xlsx(self, path: str) -> ExtractionResult:
        if HAS_PANDAS:
            df = pd.read_excel(path)
            text = df.to_string(index=False)
            tables = [{"headers": list(df.columns), "rows": df.values.tolist()}]
        else:
            text = f"[XLSX: {path} — install pandas+openpyxl for full extraction]"
            tables = []
        pages = [{"page_num": 1, "text": text, "headings": []}]
        return ExtractionResult(
            raw_text=text, pages=pages, tables=tables, headings=[],
            page_count=1,
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ocr_confidence=None, extraction_ok=True
        )


# ===========================================================================
# PHASE 4 — NORMALIZATION & CLEANING
# ===========================================================================

class Normalizer:
    """Phase 4: convert noisy extracted text into clean structured content."""

    def normalize(self, text: str, document_type: str) -> dict:
        # 1. Unicode normalisation
        text = unicodedata.normalize("NFKC", text)
        # 2. Correct OCR artifacts
        for pattern, replacement in OCR_ARTIFACTS:
            text = pattern.sub(replacement, text)
        # 3. Remove page artifacts
        text = PAGE_ARTIFACT.sub("", text)
        # 4. Collapse whitespace
        text = REPEATED_SPACE.sub(" ", text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 5. Strip leading/trailing blank lines
        lines = [l.rstrip() for l in text.split("\n")]
        text = "\n".join(lines).strip()

        # Preserve structure flags
        preserve_legal   = document_type in (DocumentType.act_regulation, DocumentType.subject_policy, DocumentType.curriculum_statement)
        preserve_curric  = document_type in (DocumentType.lesson_plan, DocumentType.textbook, DocumentType.workbook)

        curriculum_codes = CURRICULUM_CODES.findall(text) if preserve_curric else []
        legal_refs       = LEGAL_SECTION.findall(text)    if preserve_legal  else []

        language = self._detect_language(text)

        return {
            "normalized_text": text,
            "language":        language,
            "curriculum_codes": curriculum_codes,
            "legal_refs":       legal_refs,
            "word_count":       len(text.split()),
            "char_count":       len(text),
        }

    def _detect_language(self, text: str) -> str:
        """Heuristic language detection (extend with langdetect in prod)."""
        afrikaans_words = {"die", "van", "is", "en", "dat", "het", "nie", "met"}
        words = set(text.lower().split()[:200])
        if len(words & afrikaans_words) > 5:
            return "af"
        return "en"

    def infer_metadata(self, text: str, doc: Document) -> dict:
        """Phase 5: infer missing metadata from document content."""
        updates: dict = {}

        # Infer grade
        if doc.grade is None:
            m = GRADE_PATTERN.search(text[:2000])
            if m:
                g = int(m.group(1) or m.group(2))
                if 1 <= g <= 12:
                    updates["grade"] = g

        # Infer subject
        if doc.subject is None:
            text_lower = text[:3000].lower()
            for subject, keywords in SUBJECT_KEYWORDS.items():
                if any(kw in text_lower for kw in keywords):
                    updates["subject"] = subject
                    break

        # Infer publication year
        if doc.publication_year is None:
            m = re.search(r'\b(19|20)\d{2}\b', text[:1000])
            if m:
                updates["publication_year"] = int(m.group())

        # Infer title from first meaningful line
        if doc.title in ("", "Untitled"):
            first_lines = [l.strip() for l in text.split("\n") if l.strip()][:5]
            if first_lines:
                updates["title"] = first_lines[0][:120]

        # CAPS phase mapping
        if "grade" in updates or doc.grade:
            g = updates.get("grade") or doc.grade
            if g <= 3:   updates["phase"] = "Foundation Phase"
            elif g <= 6: updates["phase"] = "Intermediate Phase"
            elif g <= 9: updates["phase"] = "Senior Phase"
            else:        updates["phase"] = "FET Phase"

        return updates


# ===========================================================================
# PHASE 6 — CHUNKER
# ===========================================================================

class Chunker:
    """Phase 6: document-type-aware chunking with hierarchy preservation."""

    MAX_TOKENS = 512
    MIN_TOKENS = 30
    OVERLAP    = 50    # token overlap between adjacent chunks

    def chunk(self, text: str, document_type: str, document_id: str) -> list[DocumentChunk]:
        dispatch = {
            DocumentType.act_regulation:       self._chunk_legal,
            DocumentType.subject_policy:       self._chunk_legal,
            DocumentType.curriculum_statement: self._chunk_legal,
            DocumentType.lesson_plan:          self._chunk_lesson_plan,
            DocumentType.textbook:             self._chunk_book,
            DocumentType.workbook:             self._chunk_book,
            DocumentType.teacher_guide:        self._chunk_book,
            DocumentType.past_paper:           self._chunk_assessment,
            DocumentType.assessment_rubric:    self._chunk_assessment,
            DocumentType.memorandum:           self._chunk_assessment,
        }
        handler = dispatch.get(document_type, self._chunk_generic)
        chunks  = handler(text, document_id)
        # Assign stable IDs and timestamps
        ts = _now()
        for i, c in enumerate(chunks):
            c.chunk_id    = _uid()
            c.chunk_index = i
            c.token_count = _token_count(c.content)
            c.created_at  = ts
        return chunks

    # ── Legal: Act → Chapter → Section → Clause ──────────────────────────
    def _chunk_legal(self, text: str, document_id: str) -> list[DocumentChunk]:
        chunks = []
        section_re = re.compile(
            r'^(Chapter\s+\d+|Section\s+\d+[A-Z]?|Article\s+\d+|\d+\.\s+[A-Z])',
            re.M | re.I
        )
        parts = section_re.split(text)
        path_stack: list[str] = []
        for i in range(1, len(parts), 2):
            heading = parts[i].strip()
            body    = parts[i+1].strip() if i+1 < len(parts) else ""
            if not body:
                continue
            depth = 0 if heading.lower().startswith("chapter") else 1
            path_stack = path_stack[:depth] + [heading]
            legal_code = LEGAL_SECTION.search(heading)
            chunks.append(DocumentChunk(
                chunk_id="", document_id=document_id,
                chunk_type=ChunkType.legal_clause,
                chunk_index=0, parent_chunk_id=None,
                heading=heading, content=body,
                token_count=0,
                page_start=None, page_end=None,
                section_path=" > ".join(path_stack),
                curriculum_code=legal_code.group() if legal_code else None,
                created_at=""
            ))
        return chunks or self._chunk_generic(text, document_id)

    # ── Lesson plans: Week → Lesson → Objective → Activity ───────────────
    def _chunk_lesson_plan(self, text: str, document_id: str) -> list[DocumentChunk]:
        chunks = []
        lesson_re = re.compile(
            r'^(Week\s+\d+|Lesson\s+\d+|LESSON\s+\d+|Learning\s+Objective|Activity)',
            re.M | re.I
        )
        parts = lesson_re.split(text)
        for i in range(1, len(parts), 2):
            heading = parts[i].strip()
            body    = parts[i+1].strip() if i+1 < len(parts) else ""
            if not body:
                continue
            codes = CURRICULUM_CODES.findall(heading + " " + body[:200])
            chunks.append(DocumentChunk(
                chunk_id="", document_id=document_id,
                chunk_type=ChunkType.lesson,
                chunk_index=0, parent_chunk_id=None,
                heading=heading, content=body,
                token_count=0, page_start=None, page_end=None,
                section_path=heading,
                curriculum_code=codes[0] if codes else None,
                created_at=""
            ))
        return chunks or self._chunk_generic(text, document_id)

    # ── Books: Unit → Chapter → Topic ────────────────────────────────────
    def _chunk_book(self, text: str, document_id: str) -> list[DocumentChunk]:
        chunks = []
        header_re = re.compile(r'^(#{1,4})\s+(.+)$', re.M)
        positions = [(m.start(), len(m.group(1)), m.group(2)) for m in header_re.finditer(text)]
        positions.append((len(text), 0, ""))   # sentinel
        path_stack: list[str] = []
        for idx, (start, depth, heading) in enumerate(positions[:-1]):
            end     = positions[idx+1][0]
            body    = text[start:end].strip()
            if len(body) < 50:
                continue
            path_stack = path_stack[:depth-1] + [heading]
            codes = CURRICULUM_CODES.findall(body[:300])
            # Split oversized sections
            for sub in self._split_large(body):
                chunks.append(DocumentChunk(
                    chunk_id="", document_id=document_id,
                    chunk_type=ChunkType.topic,
                    chunk_index=0, parent_chunk_id=None,
                    heading=heading, content=sub,
                    token_count=0, page_start=None, page_end=None,
                    section_path=" > ".join(path_stack),
                    curriculum_code=codes[0] if codes else None,
                    created_at=""
                ))
        return chunks or self._chunk_generic(text, document_id)

    # ── Assessments: Question → Subquestion → Memo ───────────────────────
    def _chunk_assessment(self, text: str, document_id: str) -> list[DocumentChunk]:
        chunks = []
        q_re = re.compile(
            r'^(Question\s+\d+|Q\s*\d+\.|QUESTION\s+\d+|\d+\.\s+(?=[A-Z\(]))',
            re.M
        )
        parts = q_re.split(text)
        for i in range(1, len(parts), 2):
            heading = parts[i].strip()
            body    = parts[i+1].strip() if i+1 < len(parts) else ""
            if not body:
                continue
            chunks.append(DocumentChunk(
                chunk_id="", document_id=document_id,
                chunk_type=ChunkType.assessment_question,
                chunk_index=0, parent_chunk_id=None,
                heading=heading, content=body,
                token_count=0, page_start=None, page_end=None,
                section_path=heading, curriculum_code=None, created_at=""
            ))
        return chunks or self._chunk_generic(text, document_id)

    # ── Generic: paragraph-level chunking ────────────────────────────────
    def _chunk_generic(self, text: str, document_id: str) -> list[DocumentChunk]:
        chunks = []
        paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        buf, buf_tokens = [], 0
        for para in paragraphs:
            t = _token_count(para)
            if buf_tokens + t > self.MAX_TOKENS and buf:
                content = "\n\n".join(buf)
                chunks.append(DocumentChunk(
                    chunk_id="", document_id=document_id,
                    chunk_type=ChunkType.paragraph,
                    chunk_index=0, parent_chunk_id=None,
                    heading="", content=content,
                    token_count=0, page_start=None, page_end=None,
                    section_path="", curriculum_code=None, created_at=""
                ))
                buf, buf_tokens = [para], t
            else:
                buf.append(para); buf_tokens += t
        if buf:
            chunks.append(DocumentChunk(
                chunk_id="", document_id=document_id,
                chunk_type=ChunkType.paragraph,
                chunk_index=0, parent_chunk_id=None,
                heading="", content="\n\n".join(buf),
                token_count=0, page_start=None, page_end=None,
                section_path="", curriculum_code=None, created_at=""
            ))
        return chunks

    def _split_large(self, text: str) -> list[str]:
        """Split an oversized block into MAX_TOKENS windows."""
        words = text.split()
        step  = self.MAX_TOKENS * 4   # chars
        if len(words) * 5 <= step:
            return [text]
        results, chars = [], 0
        buf = []
        for word in words:
            buf.append(word); chars += len(word) + 1
            if chars >= step:
                results.append(" ".join(buf)); buf=[]; chars=0
        if buf: results.append(" ".join(buf))
        return results


# ===========================================================================
# PHASE 7 — QUALITY VALIDATION
# ===========================================================================

class QualityValidator:
    """Phase 7: multi-dimensional quality scoring."""

    REQUIRED_METADATA = ["title", "document_type", "language", "checksum"]
    ENRICHED_METADATA = ["subject", "grade", "curriculum", "publication_year", "license_status"]
    APPROVED_LICENSES = {LicenseStatus.open_license, LicenseStatus.creative_commons, LicenseStatus.government_open}

    def validate(self, doc: Document, chunks: list[DocumentChunk],
                 extraction: ExtractionResult) -> QualityCheckResult:
        issues: list[str] = []

        # 1. Metadata completeness (0–1)
        required_present = sum(1 for f in self.REQUIRED_METADATA if getattr(doc, f, None))
        enriched_present = sum(1 for f in self.ENRICHED_METADATA if getattr(doc, f, None))
        metadata_score = (required_present / len(self.REQUIRED_METADATA)) * 0.6 + \
                         (enriched_present / len(self.ENRICHED_METADATA)) * 0.4
        if metadata_score < 0.5:
            issues.append(f"Incomplete metadata: {required_present}/{len(self.REQUIRED_METADATA)} required fields")

        # 2. Extraction quality (0–1)
        if not extraction.extraction_ok:
            extraction_score = 0.0
            issues.append(f"Extraction failed: {extraction.error}")
        elif len(extraction.raw_text.strip()) < 100:
            extraction_score = 0.1
            issues.append("Very little text extracted — possible scanned PDF or empty document")
        else:
            text_len   = min(len(extraction.raw_text) / 5000, 1.0)
            page_bonus = min(extraction.page_count / 10, 1.0) * 0.2
            ocr_pen    = 0.0 if extraction.ocr_confidence is None else max(0, 1-extraction.ocr_confidence) * 0.3
            extraction_score = min(1.0, text_len * 0.8 + page_bonus - ocr_pen)

        # 3. Structural completeness (0–1)
        has_headings  = len(extraction.headings) > 0
        has_chunks    = len(chunks) > 0
        chunk_variety = len({c.chunk_type for c in chunks}) > 1 if chunks else False
        structure_score = (has_headings * 0.4 + has_chunks * 0.4 + chunk_variety * 0.2)
        if not has_headings:
            issues.append("No headings detected — document may lack structure")

        # 4. Completeness (0–1)
        chunk_count = len(chunks)
        if chunk_count == 0:
            completeness_score = 0.0
            issues.append("No chunks produced — document too short or extraction failed")
        elif chunk_count < 3:
            completeness_score = 0.4
            issues.append(f"Very few chunks ({chunk_count}) — document may be incomplete")
        else:
            completeness_score = min(1.0, chunk_count / 20) * 0.8 + 0.2
        # Check for duplicate chunks
        contents = [c.content[:100] for c in chunks]
        if len(contents) != len(set(contents)):
            completeness_score *= 0.8
            issues.append("Duplicate chunks detected")

        # 5. Provenance (0–1)
        has_source     = bool(doc.source_id)
        has_checksum   = bool(doc.checksum)
        has_license    = doc.license_status != LicenseStatus.unknown
        provenance_score = (has_source + has_checksum + has_license) / 3
        if not has_license:
            issues.append("License status unknown — confirm before AI use")

        # 6. Training suitability (0–1)
        license_ok   = doc.license_status in self.APPROVED_LICENSES
        quality_ok   = extraction_score > 0.5
        structure_ok = structure_score > 0.4
        training_suitability = (license_ok + quality_ok + structure_ok) / 3
        if not license_ok:
            issues.append("License not cleared for training use")

        composite = QualityCheckResult.compute_composite(
            metadata_score, extraction_score, structure_score,
            completeness_score, provenance_score, training_suitability
        )

        # Determine status
        if composite >= 0.7 and not any("failed" in i or "License" in i for i in issues):
            status = ProcessingStatus.validated
        elif composite >= 0.4:
            status = ProcessingStatus.needs_review
        else:
            status = ProcessingStatus.rejected

        return QualityCheckResult(
            document_id=doc.document_id,
            metadata_score=round(metadata_score, 3),
            extraction_score=round(extraction_score, 3),
            structure_score=round(structure_score, 3),
            completeness_score=round(completeness_score, 3),
            provenance_score=round(provenance_score, 3),
            training_suitability=round(training_suitability, 3),
            quality_score=round(composite, 3),
            status=status,
            issues=issues,
        )


# ===========================================================================
# PIPELINE ORCHESTRATOR
# ===========================================================================

class EduboostETL:
    """Full pipeline orchestrator — coordinates all phases."""

    def __init__(self, db_url: str = "sqlite:///eduboost_etl.db",
                 storage_root: str = "./data"):
        self.db_path       = db_url.replace("sqlite:///", "")
        self.storage_root  = Path(storage_root)
        self.extractor     = Extractor()
        self.normalizer    = Normalizer()
        self.chunker       = Chunker()
        self.validator     = QualityValidator()
        self._conn: Optional[sqlite3.Connection] = None

    # ── DB connection ─────────────────────────────────────────────────────
    def _db(self) -> sqlite3.Connection:
        if self._conn is None:
            self._conn = sqlite3.connect(self.db_path)
            self._conn.row_factory = sqlite3.Row
        return self._conn

    def init_db(self):
        """Create all tables. Idempotent."""
        self._db().executescript(SCHEMA_SQL)
        self._db().commit()
        self.storage_root.mkdir(parents=True, exist_ok=True)
        for sub in ("raw", "extracted", "normalized", "chunks", "rejected"):
            (self.storage_root / sub).mkdir(exist_ok=True)

    def close(self):
        if self._conn:
            self._conn.close()
            self._conn = None

    # ── Phase 1: Ingest ───────────────────────────────────────────────────
    def ingest(self, req: IngestRequest) -> Document:
        """Acquire a file, deduplicate, register source, create Document record."""
        src_path = Path(req.file_path)
        if not src_path.exists():
            raise FileNotFoundError(f"File not found: {req.file_path}")

        # Duplicate detection
        checksum = _sha256(str(src_path))
        existing = self._db().execute(
            "SELECT document_id FROM documents WHERE checksum=?", (checksum,)
        ).fetchone()
        if existing:
            raise ValueError(f"Duplicate detected: checksum {checksum[:16]}… already exists as {existing['document_id']}")

        # Register source
        source_id = _uid()
        now = _now()
        self._db().execute(
            "INSERT INTO document_sources VALUES (?,?,?,?,?,?,?,?,?)",
            (source_id, req.source_type, req.source_url, req.uploaded_by,
             req.license_status, now, checksum, 1, req.notes)
        )

        # Copy raw file to immutable store
        document_id = _uid()
        raw_dir  = self.storage_root / "raw" / source_id / document_id
        raw_dir.mkdir(parents=True, exist_ok=True)
        raw_dest = raw_dir / ("original" + src_path.suffix)
        shutil.copy2(str(src_path), str(raw_dest))

        # Create document record
        doc = Document(
            document_id=document_id,
            source_id=source_id,
            title=req.title or src_path.stem.replace("_", " ").title(),
            description="",
            document_type=req.document_type,
            subject=req.subject,
            grade=req.grade,
            phase=None, curriculum=None,
            country="ZA", province=None,
            language=req.language,
            publisher=None, author=None, publication_year=None,
            version="1.0",
            license_status=req.license_status,
            source_url=req.source_url,
            checksum=checksum,
            file_path_raw=str(raw_dest),
            file_size_bytes=src_path.stat().st_size,
            page_count=0,
            mime_type="application/octet-stream",
            processing_status=ProcessingStatus.acquired,
            quality_score=0.0,
            training_readiness=False,
            created_at=now, updated_at=now,
        )
        self._save_document(doc)
        self._db().commit()
        self._log_job(document_id, "ingest", "success")
        return doc

    # ── Phase 3: Extract ──────────────────────────────────────────────────
    def extract(self, document_id: str) -> ExtractionResult:
        doc = self._load_document(document_id)
        result = self.extractor.extract(doc.file_path_raw)

        # Persist extraction output
        ext_dir = self.storage_root / "extracted" / document_id
        ext_dir.mkdir(parents=True, exist_ok=True)
        ext_path = ext_dir / "text.json"
        ext_path.write_text(json.dumps({
            "raw_text": result.raw_text,
            "headings": result.headings,
            "page_count": result.page_count,
            "tables": result.tables,
            "extraction_ok": result.extraction_ok,
            "ocr_confidence": result.ocr_confidence,
        }, ensure_ascii=False, indent=2))

        status = ProcessingStatus.extracted if result.extraction_ok else ProcessingStatus.needs_review
        self._db().execute(
            "UPDATE documents SET processing_status=?,page_count=?,mime_type=?,updated_at=? WHERE document_id=?",
            (status, result.page_count, result.mime_type, _now(), document_id)
        )
        self._db().commit()
        self._log_job(document_id, "extract", "success" if result.extraction_ok else "failed",
                      output_path=str(ext_path))
        return result

    # ── Phase 4+5: Normalize & Enrich ─────────────────────────────────────
    def normalize(self, document_id: str) -> dict:
        doc  = self._load_document(document_id)
        ext  = self._load_extraction(document_id)
        norm = self.normalizer.normalize(ext.get("raw_text", ""), doc.document_type)

        # Phase 5: infer metadata
        updates = self.normalizer.infer_metadata(norm["normalized_text"], doc)
        updates["language"] = norm["language"]

        # Persist normalized output
        norm_dir = self.storage_root / "normalized" / document_id
        norm_dir.mkdir(parents=True, exist_ok=True)
        norm_path = norm_dir / "normalized.json"
        norm_path.write_text(json.dumps({**norm, "metadata_updates": updates}, ensure_ascii=False, indent=2))

        # Apply updates
        set_clauses = ", ".join(f"{k}=?" for k in updates)
        values = list(updates.values())
        if set_clauses:
            self._db().execute(
                f"UPDATE documents SET {set_clauses}, processing_status=?, updated_at=? WHERE document_id=?",
                [*values, ProcessingStatus.metadata_enriched, _now(), document_id]
            )
        else:
            self._db().execute(
                "UPDATE documents SET processing_status=?, updated_at=? WHERE document_id=?",
                (ProcessingStatus.metadata_enriched, _now(), document_id)
            )
        self._db().commit()
        self._log_job(document_id, "normalize", "success", str(norm_path))
        return {**norm, "metadata_updates": updates}

    # ── Phase 6: Chunk ────────────────────────────────────────────────────
    def chunk(self, document_id: str) -> list[DocumentChunk]:
        doc  = self._load_document(document_id)
        norm = self._load_normalized(document_id)
        text = norm.get("normalized_text", "")
        chunks = self.chunker.chunk(text, doc.document_type, document_id)

        # Persist chunks
        chunk_dir = self.storage_root / "chunks" / document_id
        chunk_dir.mkdir(parents=True, exist_ok=True)
        chunk_path = chunk_dir / "chunks.jsonl"
        with chunk_path.open("w") as f:
            for c in chunks:
                f.write(json.dumps(asdict(c)) + "\n")

        # Save to DB
        ts = _now()
        self._db().execute("DELETE FROM document_chunks WHERE document_id=?", (document_id,))
        for c in chunks:
            self._db().execute(
                "INSERT INTO document_chunks VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (c.chunk_id, c.document_id, c.chunk_type, c.chunk_index,
                 c.parent_chunk_id, c.heading, c.content, c.token_count,
                 c.page_start, c.page_end, c.section_path, c.curriculum_code, ts)
            )
        self._db().execute(
            "UPDATE documents SET processing_status=?, updated_at=? WHERE document_id=?",
            (ProcessingStatus.chunked, _now(), document_id)
        )
        self._db().commit()
        self._log_job(document_id, "chunk", "success", str(chunk_path))
        return chunks

    # ── Phase 7: Validate ─────────────────────────────────────────────────
    def validate(self, document_id: str) -> QualityCheckResult:
        doc    = self._load_document(document_id)
        ext    = self._load_extraction(document_id)
        chunks = self._load_chunks(document_id)

        ext_obj = ExtractionResult(
            raw_text=ext.get("raw_text",""), pages=[],
            tables=ext.get("tables",[]), headings=ext.get("headings",[]),
            page_count=ext.get("page_count",0),
            mime_type=doc.mime_type,
            ocr_confidence=ext.get("ocr_confidence"),
            extraction_ok=ext.get("extraction_ok", False),
        )

        result = self.validator.validate(doc, chunks, ext_obj)

        # Persist
        check_id = _uid()
        self._db().execute(
            "INSERT INTO quality_checks VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (check_id, document_id, result.metadata_score, result.extraction_score,
             result.structure_score, result.completeness_score, result.provenance_score,
             result.training_suitability, result.quality_score, result.status,
             json.dumps(result.issues), result.checked_at)
        )
        training_ready = result.training_suitability > 0.7 and result.quality_score > 0.7
        self._db().execute(
            "UPDATE documents SET processing_status=?, quality_score=?, training_readiness=?, updated_at=? WHERE document_id=?",
            (result.status, result.quality_score, int(training_ready), _now(), document_id)
        )
        if result.status == ProcessingStatus.needs_review:
            self._create_review_task(document_id, "; ".join(result.issues))
        self._db().commit()
        self._log_job(document_id, "validate", "success")
        return result

    # ── Full pipeline ──────────────────────────────────────────────────────
    def run_full_pipeline(self, document_id: str) -> QualityCheckResult:
        """Run all phases sequentially. Returns final quality check."""
        self.extract(document_id)
        self.normalize(document_id)
        self.chunk(document_id)
        return self.validate(document_id)

    # ── Review actions ─────────────────────────────────────────────────────
    def approve_document(self, document_id: str, reviewer: str, notes: str = "") -> Document:
        self._db().execute(
            "UPDATE documents SET processing_status=?, reviewed_by=?, reviewer_notes=?, updated_at=? WHERE document_id=?",
            (ProcessingStatus.approved, reviewer, notes, _now(), document_id)
        )
        self._db().execute(
            "UPDATE review_tasks SET resolved_at=?, resolved_by=?, resolution='approved' WHERE document_id=? AND resolved_at IS NULL",
            (_now(), reviewer, document_id)
        )
        self._db().commit()
        return self._load_document(document_id)

    def reject_document(self, document_id: str, reviewer: str, reason: str) -> Document:
        self._db().execute(
            "UPDATE documents SET processing_status=?, reviewed_by=?, rejected_reason=?, updated_at=? WHERE document_id=?",
            (ProcessingStatus.rejected, reviewer, reason, _now(), document_id)
        )
        self._db().execute(
            "UPDATE review_tasks SET resolved_at=?, resolved_by=?, resolution='rejected' WHERE document_id=? AND resolved_at IS NULL",
            (_now(), reviewer, document_id)
        )
        self._db().commit()
        return self._load_document(document_id)

    def reprocess_document(self, document_id: str) -> QualityCheckResult:
        """Reset to 'acquired' and re-run pipeline. Creates new version."""
        self._db().execute(
            "UPDATE documents SET processing_status=?, updated_at=? WHERE document_id=?",
            (ProcessingStatus.acquired, _now(), document_id)
        )
        self._db().commit()
        return self.run_full_pipeline(document_id)

    # ── Query helpers ──────────────────────────────────────────────────────
    def list_documents(self, status: Optional[str] = None,
                       grade: Optional[int] = None,
                       subject: Optional[str] = None,
                       document_type: Optional[str] = None,
                       limit: int = 100) -> list[dict]:
        clauses, params = [], []
        if status:        clauses.append("processing_status=?"); params.append(status)
        if grade:         clauses.append("grade=?");             params.append(grade)
        if subject:       clauses.append("subject=?");           params.append(subject)
        if document_type: clauses.append("document_type=?");     params.append(document_type)
        where = ("WHERE " + " AND ".join(clauses)) if clauses else ""
        rows = self._db().execute(
            f"SELECT * FROM documents {where} ORDER BY created_at DESC LIMIT ?",
            [*params, limit]
        ).fetchall()
        return [dict(r) for r in rows]

    def get_review_queue(self) -> list[dict]:
        rows = self._db().execute(
            """SELECT rt.*, d.title, d.grade, d.subject, d.document_type
               FROM review_tasks rt JOIN documents d ON rt.document_id=d.document_id
               WHERE rt.resolved_at IS NULL ORDER BY rt.created_at DESC"""
        ).fetchall()
        return [dict(r) for r in rows]

    def get_pipeline_stats(self) -> dict:
        stats = {}
        for status in ProcessingStatus:
            row = self._db().execute(
                "SELECT COUNT(*) as n FROM documents WHERE processing_status=?",
                (status.value,)
            ).fetchone()
            stats[status.value] = row["n"]
        total = self._db().execute("SELECT COUNT(*) as n FROM documents").fetchone()["n"]
        stats["total"] = total
        avg = self._db().execute("SELECT AVG(quality_score) as a FROM documents WHERE quality_score>0").fetchone()["a"]
        stats["avg_quality_score"] = round(avg or 0, 3)
        review_count = self._db().execute(
            "SELECT COUNT(*) as n FROM review_tasks WHERE resolved_at IS NULL"
        ).fetchone()["n"]
        stats["pending_reviews"] = review_count
        return stats

    def get_content_gaps(self) -> list[dict]:
        """Return combinations of grade+subject+document_type that have no approved docs."""
        rows = self._db().execute(
            """SELECT grade, subject, document_type, processing_status, COUNT(*) as cnt
               FROM documents GROUP BY grade, subject, document_type, processing_status"""
        ).fetchall()
        return [dict(r) for r in rows]

    def get_quality_report(self, document_id: str) -> dict:
        row = self._db().execute(
            "SELECT * FROM quality_checks WHERE document_id=? ORDER BY checked_at DESC LIMIT 1",
            (document_id,)
        ).fetchone()
        if not row:
            return {}
        d = dict(row)
        d["issues"] = json.loads(d.get("issues", "[]"))
        return d

    # ── Internal helpers ───────────────────────────────────────────────────
    def _save_document(self, doc: Document):
        self._db().execute(
            """INSERT OR REPLACE INTO documents VALUES
               (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (doc.document_id, doc.source_id, doc.title, doc.description,
             doc.document_type, doc.subject, doc.grade, doc.phase, doc.curriculum,
             doc.country, doc.province, doc.language, doc.publisher, doc.author,
             doc.publication_year, doc.version, doc.license_status, doc.source_url,
             doc.checksum, doc.file_path_raw, doc.file_size_bytes, doc.page_count,
             doc.mime_type, doc.processing_status, doc.quality_score,
             int(doc.training_readiness), doc.created_at, doc.updated_at,
             doc.rejected_reason, doc.reviewer_notes, doc.reviewed_by)
        )

    def _load_document(self, document_id: str) -> Document:
        row = self._db().execute(
            "SELECT * FROM documents WHERE document_id=?", (document_id,)
        ).fetchone()
        if not row:
            raise ValueError(f"Document not found: {document_id}")
        d = dict(row)
        d["training_readiness"] = bool(d.get("training_readiness", 0))
        return Document(**{k: d[k] for k in Document.__dataclass_fields__ if k in d})

    def _load_extraction(self, document_id: str) -> dict:
        path = self.storage_root / "extracted" / document_id / "text.json"
        if not path.exists():
            return {}
        return json.loads(path.read_text())

    def _load_normalized(self, document_id: str) -> dict:
        path = self.storage_root / "normalized" / document_id / "normalized.json"
        if not path.exists():
            return {}
        return json.loads(path.read_text())

    def _load_chunks(self, document_id: str) -> list[DocumentChunk]:
        rows = self._db().execute(
            "SELECT * FROM document_chunks WHERE document_id=? ORDER BY chunk_index",
            (document_id,)
        ).fetchall()
        return [DocumentChunk(**dict(r)) for r in rows]

    def _log_job(self, document_id: str, stage: str, status: str,
                 output_path: Optional[str] = None, error: Optional[str] = None):
        self._db().execute(
            "INSERT INTO processing_jobs VALUES (?,?,?,?,?,?,?,?)",
            (_uid(), document_id, stage, status, _now(), _now(), error, output_path)
        )

    def _create_review_task(self, document_id: str, reason: str):
        self._db().execute(
            "INSERT INTO review_tasks VALUES (?,?,?,?,?,?,?)",
            (_uid(), document_id, reason, _now(), None, None, None)
        )


# ===========================================================================
# QUICK-START DEMO
# ===========================================================================

if __name__ == "__main__":
    import tempfile, textwrap

    print("=== Eduboost ETL — Quick-start Demo ===\n")

    # Create a synthetic lesson plan
    sample_text = textwrap.dedent("""\
        # Grade 4 Mathematics — Lesson Plan

        ## Week 1: Introduction to Fractions

        ### Lesson 1 — What is a Fraction?

        **Learning Objective:** LO1 — Learners will identify and represent simple fractions.

        A fraction represents a part of a whole. For example, 1/2 means one out of two equal parts.

        **Activity 1:** Draw circles divided into halves, thirds, and quarters.

        ### Lesson 2 — Comparing Fractions

        **Learning Objective:** LO2 — Learners will compare fractions with the same denominator.

        When denominators are equal, the fraction with the larger numerator is greater.
        Example: 3/4 > 1/4

        ## Week 2: Addition of Fractions

        ### Lesson 3 — Adding Like Fractions

        **Learning Objective:** LO3 — Learners will add fractions with like denominators.

        To add fractions with the same denominator, add the numerators and keep the denominator.
        Example: 1/5 + 2/5 = 3/5
    """)

    with tempfile.TemporaryDirectory() as tmp:
        # Write sample file
        sample_path = Path(tmp) / "grade4_fractions_lesson_plan.md"
        sample_path.write_text(sample_text)

        # Initialise pipeline
        etl = EduboostETL(
            db_url=f"sqlite:///{tmp}/etl.db",
            storage_root=f"{tmp}/data"
        )
        etl.init_db()

        # Phase 1: Ingest
        print("Phase 1: Ingesting document...")
        doc = etl.ingest(IngestRequest(
            file_path=str(sample_path),
            source_type=SourceType.manual_upload,
            uploaded_by="content_team",
            document_type=DocumentType.lesson_plan,
            grade=4,
            subject="mathematics",
            license_status=LicenseStatus.government_open,
            title="Grade 4 Mathematics — Fractions Lesson Plan",
        ))
        print(f"  ✓ Document ID : {doc.document_id}")
        print(f"  ✓ Status      : {doc.processing_status}")
        print(f"  ✓ Checksum    : {doc.checksum[:16]}…")

        # Run full pipeline
        print("\nRunning full pipeline (Phases 3–7)...")
        result = etl.run_full_pipeline(doc.document_id)
        print(f"  ✓ Quality score   : {result.quality_score:.3f}")
        print(f"  ✓ Status          : {result.status}")
        print(f"  ✓ Issues          : {result.issues or 'none'}")

        # Inspect chunks
        chunks = etl._load_chunks(doc.document_id)
        print(f"\nChunks produced   : {len(chunks)}")
        for c in chunks[:3]:
            print(f"  [{c.chunk_type}] {c.heading[:60] or '(no heading)'!r} — {c.token_count} tokens")

        # Stats
        stats = etl.get_pipeline_stats()
        print(f"\nPipeline stats    : {stats}")

        etl.close()
        print("\nDone.")
